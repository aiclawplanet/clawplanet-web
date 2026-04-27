#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
虾蛋星球网站详细介绍文档生成脚本
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_chinese_font(run, font_name='WenQuanYi Micro Hei', size=11, bold=False):
    """设置中文字体"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if bold:
        run.font.color.rgb = RGBColor(0, 0, 0)

def add_heading_zh(doc, text, level=1):
    """添加中文标题"""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    set_chinese_font(run, size=(18 if level==1 else (14 if level==2 else 12)), bold=True)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_paragraph_zh(doc, text, bold=False, size=11):
    """添加中文段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_chinese_font(run, size=size, bold=bold)
    return p

def create_introduction_doc():
    """创建虾蛋星球详细介绍文档"""
    doc = Document()
    
    # 设置文档默认字体
    style = doc.styles['Normal']
    style.font.name = 'WenQuanYi Micro Hei'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'WenQuanYi Micro Hei')
    
    # 封面
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('虾蛋星球网站\n详细介绍')
    set_chinese_font(run, size=28, bold=True)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('\n独立开发者工具导航与需求撮合平台\n\n')
    set_chinese_font(run, size=14)
    
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run('版本：v2.0\n编制日期：2026年4月\n编制单位：虾蛋星空网络工作室\n备案号：苏ICP备2026024228号-1')
    set_chinese_font(run, size=11)
    
    doc.add_page_break()
    
    # 目录
    add_heading_zh(doc, '目录', level=1)
    toc_items = [
        '一、平台概述',
        '二、核心功能模块',
        '三、用户角色体系',
        '四、页面功能详解',
        '五、技术架构',
        '六、商业模式',
        '七、发展规划',
        '八、团队介绍',
        '九、资质与合规'
    ]
    for item in toc_items:
        add_paragraph_zh(doc, item)
    
    doc.add_page_break()
    
    # 一、平台概述
    add_heading_zh(doc, '一、平台概述', level=1)
    
    add_heading_zh(doc, '1.1 平台定位', level=2)
    add_paragraph_zh(doc, '虾蛋星球是一个专注于独立开发者生态的综合性服务平台，致力于连接独立开发者与潜在用户，同时为开发需求方提供高效的需求撮合服务。平台通过工具导航、智能推广、需求撮合三大核心业务，构建完整的独立开发者服务生态。')
    
    add_heading_zh(doc, '1.2 核心价值主张', level=2)
    values = [
        '对开发者：提供作品展示平台，获取精准流量，实现商业化变现',
        '对用户：发现优质独立开发工具，满足个性化需求',
        '对需求方：快速找到合适开发者，降低寻找成本',
        '对推广者：通过推广优质工具赚取佣金，实现流量变现'
    ]
    for v in values:
        add_paragraph_zh(doc, '• ' + v)
    
    add_heading_zh(doc, '1.3 平台数据概览', level=2)
    add_paragraph_zh(doc, '• 注册用户：持续增长中')
    add_paragraph_zh(doc, '• 入驻开发者：审核制入驻，保证质量')
    add_paragraph_zh(doc, '• 工具数量：覆盖多品类独立开发工具')
    add_paragraph_zh(doc, '• 需求撮合：一期二期功能已上线')
    
    doc.add_page_break()
    
    # 二、核心功能模块
    add_heading_zh(doc, '二、核心功能模块', level=1)
    
    add_heading_zh(doc, '2.1 工具导航平台', level=2)
    add_paragraph_zh(doc, '【功能描述】')
    add_paragraph_zh(doc, '为独立开发者提供作品展示和分发渠道，用户可以通过分类、搜索、推荐等方式发现优质工具。')
    add_paragraph_zh(doc, '【核心功能】')
    features = [
        '工具分类浏览：按效率工具、生活助手、开发工具等分类展示',
        '智能搜索：支持关键词搜索、标签筛选',
        '个性化推荐：基于用户行为的智能推荐算法',
        '工具详情页：展示工具介绍、截图、开发者信息、用户评价',
        '收藏功能：用户可收藏感兴趣的工具',
        '浏览历史：记录用户浏览过的工具'
    ]
    for f in features:
        add_paragraph_zh(doc, '• ' + f)
    
    add_heading_zh(doc, '2.2 智能推广系统', level=2)
    add_paragraph_zh(doc, '【功能描述】')
    add_paragraph_zh(doc, '为入驻开发者提供一站式推广解决方案，通过AI生成推广内容，支持多平台一键发布。')
    add_paragraph_zh(doc, '【核心功能】')
    promo_features = [
        'AI文案生成：基于工具信息自动生成推广文案',
        '多平台发布：支持微信公众号、知乎、小红书等平台',
        '推广链接生成：带追踪参数的推广链接',
        '数据统计：点击量、转化率等数据追踪',
        '星推官体系：招募推广者帮助推广，按效果付费'
    ]
    for f in promo_features:
        add_paragraph_zh(doc, '• ' + f)
    
    add_heading_zh(doc, '2.3 需求撮合系统（一期二期已上线）', level=2)
    add_paragraph_zh(doc, '【功能描述】')
    add_paragraph_zh(doc, '连接有开发需求的用户与独立开发者，提供从需求发布到对接完成的全流程服务。')
    add_paragraph_zh(doc, '【一期功能 - 信息撮合】')
    phase1 = [
        '需求发布：用户可发布网站开发、APP开发、小程序开发等需求',
        '需求大厅：开发者可浏览需求并报价',
        '报价系统：开发者提交报价，需求方选择合适开发者',
        '我的需求/报价管理：双方管理自己的需求和报价',
        '审核机制：平台审核需求，保证质量'
    ]
    for f in phase1:
        add_paragraph_zh(doc, '• ' + f)
    
    add_paragraph_zh(doc, '【二期功能 - 效率提升】')
    phase2 = [
        '即时沟通：需求方与开发者在线聊天，无需跳转第三方',
        '需求置顶：付费置顶服务，增加曝光量',
        '开发者信誉：等级体系、评分系统、历史评价',
        '星推官推广：推广需求撮合订单，赚取佣金'
    ]
    for f in phase2:
        add_paragraph_zh(doc, '• ' + f)
    
    doc.add_page_break()
    
    # 三、用户角色体系
    add_heading_zh(doc, '三、用户角色体系', level=1)
    
    add_heading_zh(doc, '3.1 普通用户', level=2)
    add_paragraph_zh(doc, '【权限功能】')
    user_funcs = [
        '浏览工具导航，搜索发现工具',
        '收藏工具，查看浏览历史',
        '发布开发需求（需手机验证）',
        '查看报价，选择开发者',
        '与开发者在线沟通',
        '评价开发者'
    ]
    for f in user_funcs:
        add_paragraph_zh(doc, '• ' + f)
    
    add_heading_zh(doc, '3.2 入驻开发者', level=2)
    add_paragraph_zh(doc, '【入驻条件】')
    add_paragraph_zh(doc, '• 提交开发者申请，通过平台审核')
    add_paragraph_zh(doc, '• 提供真实身份信息和作品案例')
    add_paragraph_zh(doc, '【权限功能】')
    dev_funcs = [
        '提交工具产品，获得展示位',
        '浏览需求大厅，查看需求详情',
        '对需求进行报价',
        '管理报价，查看报价状态',
        '与需求方在线沟通',
        '查看开发者后台数据（浏览量、点击量等）',
        '使用智能推广系统',
        '积累信誉等级和评价'
    ]
    for f in dev_funcs:
        add_paragraph_zh(doc, '• ' + f)
    
    add_heading_zh(doc, '3.3 星推官', level=2)
    add_paragraph_zh(doc, '【入驻条件】')
    add_paragraph_zh(doc, '• 提交星推官申请，通过平台审核')
    add_paragraph_zh(doc, '• 具备一定粉丝基础或推广能力')
    add_paragraph_zh(doc, '【权限功能】')
    promoter_funcs = [
        '生成工具推广链接，赚取推广佣金',
        '推广需求撮合订单，赚取撮合佣金',
        '查看推广数据统计',
        '申请佣金提现'
    ]
    for f in promoter_funcs:
        add_paragraph_zh(doc, '• ' + f)
    
    add_heading_zh(doc, '3.4 平台管理员', level=2)
    add_paragraph_zh(doc, '【权限功能】')
    admin_funcs = [
        '审核开发者入驻申请',
        '审核需求发布',
        '管理工具上下架',
        '查看平台运营数据',
        '处理用户反馈和投诉'
    ]
    for f in admin_funcs:
        add_paragraph_zh(doc, '• ' + f)
    
    doc.add_page_break()
    
    # 四、页面功能详解
    add_heading_zh(doc, '四、页面功能详解', level=1)
    
    add_heading_zh(doc, '4.1 首页', level=2)
    add_paragraph_zh(doc, '【入口路径】/')
    add_paragraph_zh(doc, '【功能模块】')
    home_modules = [
        '轮播Banner：展示平台核心功能和推广位',
        '分类导航：快速进入各工具分类',
        '搜索框：支持关键词搜索',
        '精选工具：编辑推荐优质工具',
        '最新上架：展示最新入驻的工具',
        '热门榜单：按浏览量/点击量排序',
        '个性化推荐：基于用户兴趣推荐'
    ]
    for m in home_modules:
        add_paragraph_zh(doc, '• ' + m)
    
    add_heading_zh(doc, '4.2 工具详情页', level=2)
    add_paragraph_zh(doc, '【入口路径】/tool/:id')
    add_paragraph_zh(doc, '【功能模块】')
    tool_modules = [
        '工具基本信息：名称、图标、描述、标签',
        '截图展示：工具界面截图轮播',
        '开发者信息：开发者简介、其他作品',
        '操作按钮：访问官网、收藏、分享',
        '用户评价：展示用户评论和评分',
        '相关推荐：同类工具推荐'
    ]
    for m in tool_modules:
        add_paragraph_zh(doc, '• ' + m)
    
    add_heading_zh(doc, '4.3 需求大厅', level=2)
    add_paragraph_zh(doc, '【入口路径】/demands')
    add_paragraph_zh(doc, '【功能模块】')
    demand_hall_modules = [
        '需求列表：瀑布流/列表展示所有已审核需求',
        '筛选功能：按类型、预算、周期筛选',
        '排序功能：按时间、预算排序',
        '搜索功能：关键词搜索需求',
        '发布入口：快速发布新需求'
    ]
    for m in demand_hall_modules:
        add_paragraph_zh(doc, '• ' + m)
    
    add_heading_zh(doc, '4.4 需求详情页', level=2)
    add_paragraph_zh(doc, '【入口路径】/demand/:id')
    add_paragraph_zh(doc, '【功能模块】')
    demand_detail_modules = [
        '需求信息：标题、类型、预算、周期、详细描述',
        '需求方信息：匿名展示，保护隐私',
        '报价列表（需求方可见）：查看所有报价',
        '报价按钮（开发者）：提交报价',
        '联系按钮：创建聊天会话'
    ]
    for m in demand_detail_modules:
        add_paragraph_zh(doc, '• ' + m)
    
    add_heading_zh(doc, '4.5 消息中心', level=2)
    add_paragraph_zh(doc, '【入口路径】/chat')
    add_paragraph_zh(doc, '【功能模块】')
    chat_modules = [
        '会话列表：展示所有聊天会话',
        '未读标识：显示未读消息数量',
        '聊天页面：实时消息收发、已读状态',
        '会话创建：通过需求详情页创建'
    ]
    for m in chat_modules:
        add_paragraph_zh(doc, '• ' + m)
    
    add_heading_zh(doc, '4.6 个人中心', level=2)
    add_paragraph_zh(doc, '【入口路径】/profile')
    add_paragraph_zh(doc, '【功能模块】')
    profile_modules = [
        '用户信息：头像、昵称、角色标识',
        '功能入口：我的需求、我的报价、消息中心',
        '开发者入口：开发者后台、需求大厅（开发者）',
        '星推官入口：推广中心',
        '收藏管理：我的收藏',
        '设置：账号设置、退出登录'
    ]
    for m in profile_modules:
        add_paragraph_zh(doc, '• ' + m)
    
    add_heading_zh(doc, '4.7 开发者后台', level=2)
    add_paragraph_zh(doc, '【入口路径】/dashboard')
    add_paragraph_zh(doc, '【功能模块】')
    dashboard_modules = [
        '数据统计：工具数量、总浏览量、总点击量',
        '流量趋势：近7天数据图表',
        '工具管理：添加新工具、编辑工具信息',
        '快捷入口：需求大厅、我的报价'
    ]
    for m in dashboard_modules:
        add_paragraph_zh(doc, '• ' + m)
    
    doc.add_page_break()
    
    # 五、技术架构
    add_heading_zh(doc, '五、技术架构', level=1)
    
    add_heading_zh(doc, '5.1 前端技术栈', level=2)
    add_paragraph_zh(doc, '• 框架：React 18 + TypeScript')
    add_paragraph_zh(doc, '• 构建工具：Webpack 5')
    add_paragraph_zh(doc, '• 样式：Tailwind CSS')
    add_paragraph_zh(doc, '• 路由：React Router（Hash模式）')
    add_paragraph_zh(doc, '• 动画：Framer Motion')
    add_paragraph_zh(doc, '• 图表：Recharts')
    
    add_heading_zh(doc, '5.2 后端服务', level=2)
    add_paragraph_zh(doc, '• 数据库：PostgreSQL（Supabase）')
    add_paragraph_zh(doc, '• 认证：Supabase Auth')
    add_paragraph_zh(doc, '• 存储：Supabase Storage')
    add_paragraph_zh(doc, '• 实时：Supabase Realtime')
    add_paragraph_zh(doc, '• 边缘函数：Deno TypeScript')
    
    add_heading_zh(doc, '5.3 数据库表结构', level=2)
    tables = [
        'profiles：用户资料表',
        'tools：工具信息表',
        'categories：分类表',
        'demands：需求表',
        'quotes：报价表',
        'chat_conversations：聊天会话表',
        'chat_messages：聊天消息表',
        'developer_reputation：开发者信誉表',
        'reviews：评价表',
        'promotion_links：推广链接表',
        'promoters：星推官表'
    ]
    for t in tables:
        add_paragraph_zh(doc, '• ' + t)
    
    doc.add_page_break()
    
    # 六、商业模式
    add_heading_zh(doc, '六、商业模式', level=1)
    
    add_heading_zh(doc, '6.1 收入来源', level=2)
    add_paragraph_zh(doc, '【当前已实现】')
    add_paragraph_zh(doc, '• 需求置顶服务：用户付费置顶需求，增加曝光')
    add_paragraph_zh(doc, '【规划中】')
    add_paragraph_zh(doc, '• 平台交易佣金：需求撮合成交后抽取佣金（三期）')
    add_paragraph_zh(doc, '• 会员增值服务：开发者付费会员权益')
    add_paragraph_zh(doc, '• 广告收入：首页Banner广告位')
    
    add_heading_zh(doc, '6.2 成本结构', level=2)
    costs = [
        '服务器成本：Supabase云服务费用',
        '推广成本：星推官佣金支出',
        '运营成本：平台运营、客服、审核人力成本',
        '研发成本：产品迭代、功能开发'
    ]
    for c in costs:
        add_paragraph_zh(doc, '• ' + c)
    
    doc.add_page_break()
    
    # 七、发展规划
    add_heading_zh(doc, '七、发展规划', level=1)
    
    add_heading_zh(doc, '7.1 已完成阶段', level=2)
    add_paragraph_zh(doc, '【一期 - 基础功能】')
    add_paragraph_zh(doc, '• 工具导航平台上线')
    add_paragraph_zh(doc, '• 开发者入驻系统')
    add_paragraph_zh(doc, '• 智能推广系统')
    add_paragraph_zh(doc, '• 星推官体系')
    add_paragraph_zh(doc, '【二期 - 需求撮合】')
    add_paragraph_zh(doc, '• 需求发布与审核')
    add_paragraph_zh(doc, '• 报价与对接系统')
    add_paragraph_zh(doc, '• 即时沟通功能')
    add_paragraph_zh(doc, '• 开发者信誉体系')
    add_paragraph_zh(doc, '• 需求置顶服务')
    
    add_heading_zh(doc, '7.2 未来规划', level=2)
    add_paragraph_zh(doc, '【三期 - 商业化】（需ICP经营许可证）')
    add_paragraph_zh(doc, '• 平台担保交易')
    add_paragraph_zh(doc, '• 资金托管功能')
    add_paragraph_zh(doc, '• 里程碑付款')
    add_paragraph_zh(doc, '• 平台交易抽佣')
    
    add_paragraph_zh(doc, '【长期规划】')
    add_paragraph_zh(doc, '• 开发者社区建设')
    add_paragraph_zh(doc, '• 工具API市场')
    add_paragraph_zh(doc, '• 海外市场拓展')
    
    doc.add_page_break()
    
    # 八、团队介绍
    add_heading_zh(doc, '八、团队介绍', level=1)
    add_paragraph_zh(doc, '虾蛋星空网络工作室成立于2024年，是一家专注于独立开发者生态的互联网创业团队。')
    add_paragraph_zh(doc, '团队核心成员拥有丰富的互联网产品开发经验，致力于通过技术创新帮助独立开发者实现商业价值。')
    
    # 九、资质与合规
    add_heading_zh(doc, '九、资质与合规', level=1)
    
    add_heading_zh(doc, '9.1 现有资质', level=2)
    add_paragraph_zh(doc, '• 营业执照：南京市玄武区虾蛋星空网络工作室（个体工商户）')
    add_paragraph_zh(doc, '• ICP备案：苏ICP备2026024228号-1')
    add_paragraph_zh(doc, '• 域名：aiclawplanet.com')
    
    add_heading_zh(doc, '9.2 合规说明', level=2)
    add_paragraph_zh(doc, '【一期二期合规性】')
    add_paragraph_zh(doc, '• 信息撮合服务属于"信息技术咨询服务"，在经营范围内')
    add_paragraph_zh(doc, '• 仅提供信息展示和对接，不涉及资金流转')
    add_paragraph_zh(doc, '• 平台已展示风险提示文案，明确平台责任边界')
    add_paragraph_zh(doc, '【三期资质要求】')
    add_paragraph_zh(doc, '• 资金托管功能需办理ICP经营许可证')
    add_paragraph_zh(doc, '• 建议三期前注册公司，办理相关资质')
    
    # 保存文档
    doc.save('/home/project/outputs/虾蛋星球网站详细介绍_v2.0.docx')
    print("详细介绍文档已生成：/home/project/outputs/虾蛋星球网站详细介绍_v2.0.docx")

if __name__ == '__main__':
    create_introduction_doc()
