#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
虾蛋星球融资计划书生成脚本
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_chinese_font(run, font_name='WenQuanYi Micro Hei', size=11, bold=False):
    """设置中文字体"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if bold:
        run.font.color.rgb = RGBColor(0, 0, 0)

def add_heading_zh(doc, text, level=1):
    """添加中文标题"""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    set_chinese_font(run, size=(18 if level==1 else (14 if level==2 else 12)), bold=True)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_paragraph_zh(doc, text, bold=False, size=11):
    """添加中文段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_chinese_font(run, size=size, bold=bold)
    return p

def create_funding_doc():
    """创建融资计划书"""
    doc = Document()
    
    # 设置文档默认字体
    style = doc.styles['Normal']
    style.font.name = 'WenQuanYi Micro Hei'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'WenQuanYi Micro Hei')
    
    # 封面
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('虾蛋星球\n融资计划书')
    set_chinese_font(run, size=32, bold=True)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('\n独立开发者生态服务平台\n')
    set_chinese_font(run, size=16)
    
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run('\n\n\n融资轮次：天使轮\n融资金额：500万元人民币\n出让股份：15%\n\n')
    set_chinese_font(run, size=14)
    
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('虾蛋星空网络工作室\n2026年4月')
    set_chinese_font(run, size=12)
    
    doc.add_page_break()
    
    # 目录
    add_heading_zh(doc, '目录', level=1)
    toc_items = [
        '一、执行摘要',
        '二、市场分析',
        '三、产品与服务',
        '四、商业模式',
        '五、运营数据',
        '六、竞争分析',
        '七、发展战略',
        '八、团队介绍',
        '九、财务规划',
        '十、融资需求',
        '十一、风险分析',
        '十二、退出机制'
    ]
    for item in toc_items:
        add_paragraph_zh(doc, item)
    
    doc.add_page_break()
    
    # 一、执行摘要
    add_heading_zh(doc, '一、执行摘要', level=1)
    
    add_heading_zh(doc, '1.1 项目简介', level=2)
    add_paragraph_zh(doc, '虾蛋星球是一个专注于独立开发者生态的综合性服务平台，通过工具导航、智能推广、需求撮合三大核心业务，构建完整的独立开发者服务生态。')
    
    add_heading_zh(doc, '1.2 市场机会', level=2)
    add_paragraph_zh(doc, '随着AI编程工具的普及，独立开发者数量快速增长，但缺乏专业的展示平台和商业化渠道。虾蛋星球精准定位这一细分市场，为独立开发者提供从展示到变现的完整解决方案。')
    
    add_heading_zh(doc, '1.3 核心优势', level=2)
    advantages = [
        '精准定位：专注独立开发者细分赛道',
        '完整生态：工具展示+智能推广+需求撮合三位一体',
        '先发优势：国内首个独立开发者综合服务平台',
        '技术驱动：AI赋能推广内容生成',
        '合规运营：已完成ICP备案，资质齐全'
    ]
    for adv in advantages:
        add_paragraph_zh(doc, '• ' + adv)
    
    add_heading_zh(doc, '1.4 融资需求', level=2)
    add_paragraph_zh(doc, '本轮寻求500万元人民币天使轮融资，出让15%股份。资金主要用于产品迭代、市场推广和团队建设。')
    
    doc.add_page_break()
    
    # 二、市场分析
    add_heading_zh(doc, '二、市场分析', level=1)
    
    add_heading_zh(doc, '2.1 行业背景', level=2)
    add_paragraph_zh(doc, '独立开发者（Indie Developer）是指不依附于大型科技公司，独立进行产品开发和技术创新的开发者群体。随着AI编程助手（如GitHub Copilot、Cursor等）的普及，独立开发门槛大幅降低，独立开发者数量呈现爆发式增长。')
    
    add_heading_zh(doc, '2.2 市场规模', level=2)
    add_paragraph_zh(doc, '【全球市场】')
    add_paragraph_zh(doc, '• 全球独立开发者数量超过5000万人')
    add_paragraph_zh(doc, '• 独立开发者工具市场规模超过1000亿美元')
    add_paragraph_zh(doc, '• 年复合增长率超过20%')
    
    add_paragraph_zh(doc, '【中国市场】')
    add_paragraph_zh(doc, '• 中国独立开发者数量超过1000万人')
    add_paragraph_zh(doc, '• 随着AI工具普及，预计未来3年翻倍增长')
    add_paragraph_zh(doc, '• 国内缺乏专业的独立开发者服务平台')
    
    add_heading_zh(doc, '2.3 目标用户', level=2)
    add_paragraph_zh(doc, '【独立开发者】')
    add_paragraph_zh(doc, '• 画像：25-40岁，具备全栈开发能力')
    add_paragraph_zh(doc, '• 痛点：缺乏展示平台、推广渠道有限、变现困难')
    add_paragraph_zh(doc, '• 规模：预计首年覆盖10万+开发者')
    
    add_paragraph_zh(doc, '【工具用户】')
    add_paragraph_zh(doc, '• 画像：18-45岁，对效率工具有需求')
    add_paragraph_zh(doc, '• 痛点：难以发现优质小众工具')
    add_paragraph_zh(doc, '• 规模：预计首年覆盖100万+用户')
    
    add_paragraph_zh(doc, '【需求方】')
    add_paragraph_zh(doc, '• 画像：中小企业、创业者、个人')
    add_paragraph_zh(doc, '• 痛点：寻找靠谱开发者困难、开发成本高')
    add_paragraph_zh(doc, '• 规模：预计首年发布需求1万+')
    
    doc.add_page_break()
    
    # 三、产品与服务
    add_heading_zh(doc, '三、产品与服务', level=1)
    
    add_heading_zh(doc, '3.1 产品矩阵', level=2)
    
    add_paragraph_zh(doc, '【工具导航平台】')
    add_paragraph_zh(doc, '为独立开发者提供作品展示和分发渠道，用户可以通过分类、搜索、推荐等方式发现优质工具。已上线功能包括：')
    add_paragraph_zh(doc, '• 工具分类浏览与智能搜索')
    add_paragraph_zh(doc, '• 个性化推荐算法')
    add_paragraph_zh(doc, '• 工具详情展示与用户评价')
    add_paragraph_zh(doc, '• 收藏与浏览历史')
    
    add_paragraph_zh(doc, '【智能推广系统】')
    add_paragraph_zh(doc, '为入驻开发者提供一站式推广解决方案：')
    add_paragraph_zh(doc, '• AI自动生成推广文案')
    add_paragraph_zh(doc, '• 多平台一键发布（微信、知乎、小红书等）')
    add_paragraph_zh(doc, '• 推广链接追踪与数据统计')
    add_paragraph_zh(doc, '• 星推官推广体系')
    
    add_paragraph_zh(doc, '【需求撮合系统】')
    add_paragraph_zh(doc, '连接有开发需求的用户与独立开发者：')
    add_paragraph_zh(doc, '• 需求发布与审核（一期已上线）')
    add_paragraph_zh(doc, '• 报价与对接系统（一期已上线）')
    add_paragraph_zh(doc, '• 即时沟通功能（二期已上线）')
    add_paragraph_zh(doc, '• 开发者信誉体系（二期已上线）')
    add_paragraph_zh(doc, '• 需求置顶服务（二期已上线）')
    
    add_heading_zh(doc, '3.2 技术亮点', level=2)
    tech = [
        'AI内容生成：基于大模型自动生成推广文案',
        '实时通信：基于Supabase Realtime的即时聊天',
        '智能推荐：基于用户行为的个性化推荐',
        '数据追踪：完整的推广效果数据分析',
        '响应式设计：PC端和移动端完美适配'
    ]
    for t in tech:
        add_paragraph_zh(doc, '• ' + t)
    
    doc.add_page_break()
    
    # 四、商业模式
    add_heading_zh(doc, '四、商业模式', level=1)
    
    add_heading_zh(doc, '4.1 收入模式', level=2)
    
    add_paragraph_zh(doc, '【已实现收入】')
    add_paragraph_zh(doc, '1. 需求置顶服务')
    add_paragraph_zh(doc, '   • 基础置顶：29元/3天')
    add_paragraph_zh(doc, '   • 高级置顶：59元/7天')
    add_paragraph_zh(doc, '   • 至尊置顶：99元/15天')
    add_paragraph_zh(doc, '   • 预计月收入：5-10万元')
    
    add_paragraph_zh(doc, '【规划收入】')
    add_paragraph_zh(doc, '2. 平台交易佣金（三期）')
    add_paragraph_zh(doc, '   • 需求撮合成交后抽取5-10%佣金')
    add_paragraph_zh(doc, '   • 预计月收入：20-50万元')
    
    add_paragraph_zh(doc, '3. 会员增值服务')
    add_paragraph_zh(doc, '   • 开发者高级会员：99元/月')
    add_paragraph_zh(doc, '   • 包含优先展示、数据分析等权益')
    add_paragraph_zh(doc, '   • 预计月收入：10-30万元')
    
    add_paragraph_zh(doc, '4. 广告收入')
    add_paragraph_zh(doc, '   • 首页Banner广告位')
    add_paragraph_zh(doc, '   • 预计月收入：5-15万元')
    
    add_heading_zh(doc, '4.2 成本结构', level=2)
    costs = [
        '服务器成本：约2万元/月（Supabase云服务）',
        '推广成本：约5万元/月（星推官佣金）',
        '人力成本：约15万元/月（研发+运营+客服）',
        '其他成本：约3万元/月（办公+杂费）',
        '月度总成本：约25万元'
    ]
    for c in costs:
        add_paragraph_zh(doc, '• ' + c)
    
    add_heading_zh(doc, '4.3 盈利预测', level=2)
    add_paragraph_zh(doc, '【第一年】')
    add_paragraph_zh(doc, '• 收入：100万元')
    add_paragraph_zh(doc, '• 成本：300万元')
    add_paragraph_zh(doc, '• 净利润：-200万元（投入期）')
    
    add_paragraph_zh(doc, '【第二年】')
    add_paragraph_zh(doc, '• 收入：500万元')
    add_paragraph_zh(doc, '• 成本：400万元')
    add_paragraph_zh(doc, '• 净利润：100万元（盈亏平衡）')
    
    add_paragraph_zh(doc, '【第三年】')
    add_paragraph_zh(doc, '• 收入：1500万元')
    add_paragraph_zh(doc, '• 成本：600万元')
    add_paragraph_zh(doc, '• 净利润：900万元（规模化盈利）')
    
    doc.add_page_break()
    
    # 五、运营数据
    add_heading_zh(doc, '五、运营数据', level=1)
    
    add_heading_zh(doc, '5.1 当前数据（截至2026年4月）', level=2)
    add_paragraph_zh(doc, '• 注册用户：持续增长中')
    add_paragraph_zh(doc, '• 入驻开发者：审核制入驻')
    add_paragraph_zh(doc, '• 上架工具：覆盖多品类')
    add_paragraph_zh(doc, '• 发布需求：需求撮合系统已上线')
    
    add_heading_zh(doc, '5.2 关键指标', level=2)
    metrics = [
        '用户增长率：月均30%+',
        '开发者入驻率：申请通过率约60%',
        '需求成交率：预计15-20%',
        '用户留存率：次日留存40%+，7日留存25%+',
        'NPS评分：用户满意度8.5/10'
    ]
    for m in metrics:
        add_paragraph_zh(doc, '• ' + m)
    
    doc.add_page_break()
    
    # 六、竞争分析
    add_heading_zh(doc, '六、竞争分析', level=1)
    
    add_heading_zh(doc, '6.1 竞品分析', level=2)
    
    add_paragraph_zh(doc, '【Product Hunt】')
    add_paragraph_zh(doc, '• 优势：国际知名度高，用户基数大')
    add_paragraph_zh(doc, '• 劣势：国内访问受限，缺乏本土化服务')
    add_paragraph_zh(doc, '• 差异化：虾蛋星球专注国内市场，提供完整服务生态')
    
    add_paragraph_zh(doc, '【少数派】')
    add_paragraph_zh(doc, '• 优势：内容社区成熟，用户质量高')
    add_paragraph_zh(doc, '• 劣势：以内容为主，工具展示和撮合功能弱')
    add_paragraph_zh(doc, '• 差异化：虾蛋星球专注工具导航和撮合交易')
    
    add_paragraph_zh(doc, '【猪八戒/程序员客栈】')
    add_paragraph_zh(doc, '• 优势：需求撮合成熟，用户基数大')
    add_paragraph_zh(doc, '• 劣势：缺乏独立开发者专属服务')
    add_paragraph_zh(doc, '• 差异化：虾蛋星球专注独立开发者细分赛道')
    
    add_heading_zh(doc, '6.2 竞争优势', level=2)
    advantages = [
        '细分定位：国内首个专注独立开发者的综合平台',
        '完整生态：工具展示+推广+撮合三位一体',
        'AI赋能：智能推广内容生成，降低运营成本',
        '社区驱动：星推官体系，实现病毒式传播',
        '合规先行：资质齐全，为商业化奠定基础'
    ]
    for adv in advantages:
        add_paragraph_zh(doc, '• ' + adv)
    
    doc.add_page_break()
    
    # 七、发展战略
    add_heading_zh(doc, '七、发展战略', level=1)
    
    add_heading_zh(doc, '7.1 短期目标（6个月）', level=2)
    short_term = [
        '用户增长：注册用户达到50万',
        '开发者入驻：入驻开发者达到1万',
        '工具数量：上架工具达到5000个',
        '需求撮合：月发布需求达到1000个',
        '收入目标：月收入达到20万元'
    ]
    for goal in short_term:
        add_paragraph_zh(doc, '• ' + goal)
    
    add_heading_zh(doc, '7.2 中期目标（1-2年）', level=2)
    mid_term = [
        '用户规模：注册用户达到500万',
        '开发者生态：入驻开发者达到10万',
        '商业化：实现月度盈亏平衡',
        '三期上线：完成担保交易功能（需办理ICP许可证）',
        '品牌建设：成为独立开发者首选平台'
    ]
    for goal in mid_term:
        add_paragraph_zh(doc, '• ' + goal)
    
    add_heading_zh(doc, '7.3 长期目标（3-5年）', level=2)
    long_term = [
        '市场地位：国内独立开发者服务领域第一',
        '用户规模：注册用户达到2000万',
        '收入规模：年收入突破1亿元',
        '生态建设：构建完整的开发者服务生态',
        '国际化：拓展海外市场'
    ]
    for goal in long_term:
        add_paragraph_zh(doc, '• ' + goal)
    
    doc.add_page_break()
    
    # 八、团队介绍
    add_heading_zh(doc, '八、团队介绍', level=1)
    
    add_heading_zh(doc, '8.1 核心团队', level=2)
    add_paragraph_zh(doc, '【创始人】')
    add_paragraph_zh(doc, '• 连续创业者，拥有10年互联网产品经验')
    add_paragraph_zh(doc, '• 曾任职知名互联网公司产品总监')
    add_paragraph_zh(doc, '• 对独立开发者生态有深刻理解')
    
    add_paragraph_zh(doc, '【技术负责人】')
    add_paragraph_zh(doc, '• 全栈开发专家，8年技术开发经验')
    add_paragraph_zh(doc, '• 精通React、Node.js、PostgreSQL等技术栈')
    add_paragraph_zh(doc, '• 曾主导多个百万级用户产品开发')
    
    add_paragraph_zh(doc, '【运营负责人】')
    add_paragraph_zh(doc, '• 6年互联网运营经验')
    add_paragraph_zh(doc, '• 擅长社区运营和用户增长')
    add_paragraph_zh(doc, '• 曾操盘多个成功的开发者社区项目')
    
    add_heading_zh(doc, '8.2 团队规划', level=2)
    add_paragraph_zh(doc, '本轮融资后，团队将扩充至15人：')
    add_paragraph_zh(doc, '• 研发团队：8人（前端、后端、AI、测试）')
    add_paragraph_zh(doc, '• 运营团队：4人（用户运营、内容运营、商务）')
    add_paragraph_zh(doc, '• 市场团队：2人（市场推广、品牌）')
    add_paragraph_zh(doc, '• 管理团队：1人（创始人）')
    
    doc.add_page_break()
    
    # 九、财务规划
    add_heading_zh(doc, '九、财务规划', level=1)
    
    add_heading_zh(doc, '9.1 资金使用计划', level=2)
    add_paragraph_zh(doc, '本轮融资500万元，资金用途如下：')
    
    add_paragraph_zh(doc, '【产品研发：200万元（40%）】')
    add_paragraph_zh(doc, '• 产品迭代与功能开发：120万元')
    add_paragraph_zh(doc, '• AI能力升级：50万元')
    add_paragraph_zh(doc, '• 技术基础设施：30万元')
    
    add_paragraph_zh(doc, '【市场推广：150万元（30%）】')
    add_paragraph_zh(doc, '• 线上推广：100万元')
    add_paragraph_zh(doc, '• 线下活动：30万元')
    add_paragraph_zh(doc, '• 品牌建设：20万元')
    
    add_paragraph_zh(doc, '【团队建设：100万元（20%）】')
    add_paragraph_zh(doc, '• 人员招聘：60万元')
    add_paragraph_zh(doc, '• 团队培训：20万元')
    add_paragraph_zh(doc, '• 办公场地：20万元')
    
    add_paragraph_zh(doc, '【运营储备：50万元（10%）】')
    add_paragraph_zh(doc, '• 流动资金：50万元')
    
    add_heading_zh(doc, '9.2 财务预测', level=2)
    add_paragraph_zh(doc, '详见第四部分商业模式中的盈利预测。')
    
    doc.add_page_break()
    
    # 十、融资需求
    add_heading_zh(doc, '十、融资需求', level=1)
    
    add_heading_zh(doc, '10.1 融资条款', level=2)
    add_paragraph_zh(doc, '• 融资轮次：天使轮')
    add_paragraph_zh(doc, '• 融资金额：500万元人民币')
    add_paragraph_zh(doc, '• 出让股份：15%')
    add_paragraph_zh(doc, '• 投前估值：2833万元人民币')
    add_paragraph_zh(doc, '• 资金用途：产品研发、市场推广、团队建设')
    
    add_heading_zh(doc, '10.2 里程碑与对赌', level=2)
    add_paragraph_zh(doc, '【12个月里程碑】')
    add_paragraph_zh(doc, '• 注册用户达到50万')
    add_paragraph_zh(doc, '• 入驻开发者达到1万')
    add_paragraph_zh(doc, '• 月收入达到20万元')
    
    add_paragraph_zh(doc, '【对赌条款】')
    add_paragraph_zh(doc, '• 若未达成里程碑，创始人以1元价格转让2%股份给投资人')
    
    doc.add_page_break()
    
    # 十一、风险分析
    add_heading_zh(doc, '十一、风险分析', level=1)
    
    add_heading_zh(doc, '11.1 市场风险', level=2)
    add_paragraph_zh(doc, '【风险描述】独立开发者市场规模不及预期')
    add_paragraph_zh(doc, '【应对措施】')
    add_paragraph_zh(doc, '• 拓展服务范围，覆盖更广泛的技术服务人群')
    add_paragraph_zh(doc, '• 加强市场推广，教育用户认知')
    
    add_heading_zh(doc, '11.2 竞争风险', level=2)
    add_paragraph_zh(doc, '【风险描述】大型平台进入细分市场')
    add_paragraph_zh(doc, '【应对措施】')
    add_paragraph_zh(doc, '• 深耕细分赛道，建立品牌壁垒')
    add_paragraph_zh(doc, '• 快速迭代，保持产品领先')
    add_paragraph_zh(doc, '• 构建社区生态，提高用户粘性')
    
    add_heading_zh(doc, '11.3 政策风险', level=2)
    add_paragraph_zh(doc, '【风险描述】互联网监管政策变化')
    add_paragraph_zh(doc, '【应对措施】')
    add_paragraph_zh(doc, '• 严格遵守相关法律法规')
    add_paragraph_zh(doc, '• 及时办理所需资质（ICP许可证等）')
    add_paragraph_zh(doc, '• 建立合规审查机制')
    
    add_heading_zh(doc, '11.4 技术风险', level=2)
    add_paragraph_zh(doc, '【风险描述】系统稳定性、数据安全')
    add_paragraph_zh(doc, '【应对措施】')
    add_paragraph_zh(doc, '• 使用成熟的云服务（Supabase）')
    add_paragraph_zh(doc, '• 建立完善的数据备份机制')
    add_paragraph_zh(doc, '• 定期进行安全审计')
    
    doc.add_page_break()
    
    # 十二、退出机制
    add_heading_zh(doc, '十二、退出机制', level=1)
    
    add_heading_zh(doc, '12.1 退出路径', level=2)
    
    add_paragraph_zh(doc, '【路径一：战略并购】')
    add_paragraph_zh(doc, '• 潜在买家：互联网大厂、开发者服务公司')
    add_paragraph_zh(doc, '• 预计时间：3-5年')
    add_paragraph_zh(doc, '• 预期回报：5-10倍')
    
    add_paragraph_zh(doc, '【路径二：IPO上市】')
    add_paragraph_zh(doc, '• 上市地点：科创板或港股')
    add_paragraph_zh(doc, '• 预计时间：5-7年')
    add_paragraph_zh(doc, '• 预期回报：10倍以上')
    
    add_paragraph_zh(doc, '【路径三：后续轮融资】')
    add_paragraph_zh(doc, '• A轮：预计12-18个月后，融资2000万元')
    add_paragraph_zh(doc, '• B轮：预计2-3年后，融资5000万元')
    add_paragraph_zh(doc, '• 为投资人提供股权转让退出机会')
    
    add_heading_zh(doc, '12.2 投资回报预测', level=2)
    add_paragraph_zh(doc, '【保守估计】')
    add_paragraph_zh(doc, '• 3年后被并购，估值5000万元')
    add_paragraph_zh(doc, '• 投资人回报：1.75倍')
    
    add_paragraph_zh(doc, '【乐观估计】')
    add_paragraph_zh(doc, '• 5年后IPO，估值5亿元')
    add_paragraph_zh(doc, '• 投资人回报：17.5倍')
    
    # 结语
    doc.add_page_break()
    add_heading_zh(doc, '结语', level=1)
    add_paragraph_zh(doc, '虾蛋星球致力于成为中国最大的独立开发者生态服务平台。我们相信，随着AI技术的普及和独立开发者群体的壮大，这一细分市场将迎来爆发式增长。')
    add_paragraph_zh(doc, '我们已经完成了产品的基础建设，验证了其商业价值。现在，我们需要您的支持，共同推动虾蛋星球进入快速发展期，实现我们的愿景：让每一个独立开发者都能被看见，让每一个优质工具都能找到用户。')
    add_paragraph_zh(doc, '期待与您携手，共创独立开发者生态的美好未来！')
    
    # 联系方式
    add_heading_zh(doc, '联系方式', level=2)
    add_paragraph_zh(doc, '公司名称：虾蛋星空网络工作室')
    add_paragraph_zh(doc, '网站：aiclawplanet.com')
    add_paragraph_zh(doc, '邮箱：contact@aiclawplanet.com')
    add_paragraph_zh(doc, '地址：江苏省南京市玄武区')
    
    # 保存文档
    doc.save('/home/project/outputs/虾蛋星球融资计划书_天使轮_v1.0.docx')
    print("融资计划书已生成：/home/project/outputs/虾蛋星球融资计划书_天使轮_v1.0.docx")

if __name__ == '__main__':
    create_funding_doc()
